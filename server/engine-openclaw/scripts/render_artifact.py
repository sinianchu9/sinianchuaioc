#!/usr/bin/env python3
import argparse
import base64
import csv
import hashlib
import json
import os
from pathlib import Path
from typing import Any


def _safe_name(name: str) -> str:
    bad = '<>:"/\\|?*'
    cleaned = "".join("_" if c in bad or ord(c) < 32 else c for c in name).strip()
    return cleaned or f"artifact-{os.getpid()}.json"


def _write_text(path: Path, text: str) -> None:
    path.write_text(text, encoding="utf-8")


def _render_docx(path: Path, payload: dict[str, Any]) -> tuple[str, str]:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return "fallback", "python-docx unavailable"

    doc = Document()
    title = str(payload.get("title") or path.stem)
    doc.add_heading(title, level=1)
    summary = payload.get("summary")
    if summary:
        doc.add_paragraph(str(summary))

    sections = payload.get("sections")
    if isinstance(sections, list):
        for sec in sections:
            if not isinstance(sec, dict):
                continue
            heading = str(sec.get("h") or sec.get("title") or "Section")
            doc.add_heading(heading, level=2)
            bullets = sec.get("bullets") or sec.get("items") or []
            if isinstance(bullets, list):
                for b in bullets:
                    doc.add_paragraph(str(b), style="List Bullet")

    doc.save(path)
    return "docx", ""


def _render_xlsx(path: Path, payload: dict[str, Any]) -> tuple[str, str]:
    try:
        from openpyxl import Workbook  # type: ignore
    except Exception:
        return "fallback", "openpyxl unavailable"

    wb = Workbook()
    wb.remove(wb.active)
    sheets = payload.get("sheets")
    if isinstance(sheets, list) and sheets:
        for sheet in sheets:
            if not isinstance(sheet, dict):
                continue
            ws = wb.create_sheet(str(sheet.get("name") or "sheet"))
            columns = sheet.get("columns") or []
            if isinstance(columns, list) and columns:
                ws.append([str(c) for c in columns])
            rows = sheet.get("rows") or []
            if isinstance(rows, list):
                for row in rows:
                    if isinstance(row, list):
                        ws.append([str(x) for x in row])
    else:
        ws = wb.create_sheet("sheet1")
        ws.append(["key", "value"])
        for k, v in payload.items():
            ws.append([str(k), json.dumps(v, ensure_ascii=False)])
    wb.save(path)
    return "xlsx", ""


def _render_pdf(path: Path, payload: dict[str, Any]) -> tuple[str, str]:
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
    except Exception:
        return "fallback", "reportlab unavailable"

    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(40, y, str(payload.get("title") or path.stem))
    y -= 24
    c.setFont("Helvetica", 10)
    body = payload.get("summary") or payload.get("markdown") or json.dumps(payload, ensure_ascii=False)
    for line in str(body).splitlines():
        if y < 40:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = height - 50
        c.drawString(40, y, line[:120])
        y -= 14
    c.save()
    return "pdf", ""


def _render_csv(path: Path, payload: dict[str, Any]) -> tuple[str, str]:
    columns = payload.get("columns") if isinstance(payload, dict) else []
    rows = payload.get("rows") if isinstance(payload, dict) else []
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        if isinstance(columns, list):
            w.writerow([str(c) for c in columns])
        if isinstance(rows, list):
            for row in rows:
                if isinstance(row, list):
                    w.writerow([str(x) for x in row])
    return "csv", ""


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--type", required=True)
    ap.add_argument("--filename", required=True)
    ap.add_argument("--dir", required=True)
    ap.add_argument("--payload-b64", required=True)
    args = ap.parse_args()

    out_dir = Path(args.dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = _safe_name(args.filename)
    out_path = out_dir / filename
    payload_raw = base64.b64decode(args.payload_b64.encode("utf-8")).decode("utf-8", errors="replace")
    payload = json.loads(payload_raw) if payload_raw.strip() else {}
    if not isinstance(payload, dict):
        payload = {"value": payload}

    out_type = args.type.lower()
    actual_type = out_type
    warning = ""

    if out_type in {"json"}:
        _write_text(out_path, json.dumps(payload, ensure_ascii=False, indent=2))
    elif out_type in {"md", "markdown"}:
        _write_text(out_path, str(payload.get("markdown") or payload.get("summary") or json.dumps(payload, ensure_ascii=False, indent=2)))
    elif out_type in {"txt"}:
        _write_text(out_path, str(payload.get("text") or payload.get("summary") or json.dumps(payload, ensure_ascii=False, indent=2)))
    elif out_type in {"csv"}:
        actual_type, warning = _render_csv(out_path, payload)
    elif out_type in {"docx"}:
        actual_type, warning = _render_docx(out_path, payload)
    elif out_type in {"xlsx"}:
        actual_type, warning = _render_xlsx(out_path, payload)
    elif out_type in {"pdf"}:
        actual_type, warning = _render_pdf(out_path, payload)
    else:
        _write_text(out_path, json.dumps(payload, ensure_ascii=False, indent=2))
        warning = f"unsupported type {out_type}, wrote json payload"
        actual_type = "json"

    if actual_type == "fallback":
        fallback_path = out_path.with_suffix(".json")
        _write_text(
            fallback_path,
            json.dumps(
                {
                    "requested_type": out_type,
                    "warning": warning,
                    "payload": payload,
                },
                ensure_ascii=False,
                indent=2,
            ),
        )
        out_path = fallback_path
        actual_type = "json"

    raw = out_path.read_bytes()
    out = {
        "output_path": str(out_path),
        "filename": out_path.name,
        "actual_type": actual_type,
        "size_bytes": len(raw),
        "sha256": hashlib.sha256(raw).hexdigest(),
        "warning": warning,
    }
    print(json.dumps(out, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
